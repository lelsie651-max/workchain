from __future__ import annotations

from io import BytesIO
import time

from docx import Document
from pypdf import PdfReader

from app import ocr


MAX_EXTRACTED_TEXT_LENGTH = 50_000


def _detect_image_mime_type(payload: bytes, filename: str) -> str:
    lower_name = (filename or "").lower()
    if payload.startswith(b"\x89PNG\r\n\x1a\n") or lower_name.endswith(".png"):
        return "image/png"
    if payload.startswith(b"\xff\xd8\xff") or lower_name.endswith((".jpg", ".jpeg")):
        return "image/jpeg"
    if payload.startswith((b"GIF87a", b"GIF89a")) or lower_name.endswith(".gif"):
        return "image/gif"
    if (len(payload) >= 12 and payload.startswith(b"RIFF") and payload[8:12] == b"WEBP") or lower_name.endswith(".webp"):
        return "image/webp"
    return "image/jpeg"


def _limit_text(text: str) -> tuple[str, str]:
    if len(text) <= MAX_EXTRACTED_TEXT_LENGTH:
        return text, ""
    return text[:MAX_EXTRACTED_TEXT_LENGTH], f"已截取前 {MAX_EXTRACTED_TEXT_LENGTH} 个字符"


def _extract_pdf_text(payload: bytes) -> tuple[str | None, str]:
    try:
        reader = PdfReader(BytesIO(payload))
        page_texts = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as exc:
        return None, f"这份 PDF 暂时无法读取({type(exc).__name__})"

    text = "\n\n".join(item for item in page_texts if item).strip()
    if not text:
        return None, "这份 PDF 看起来是扫描件,没有可提取的文字"
    return _limit_text(text)


def _extract_docx_text(payload: bytes) -> tuple[str | None, str]:
    try:
        document = Document(BytesIO(payload))
    except Exception as exc:
        return None, f"这份 docx 暂时无法读取({type(exc).__name__})"

    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        return None, "这份 docx 里没有可提取的文字"
    return _limit_text(text)


def _extract_txt_text(payload: bytes) -> tuple[str | None, str]:
    for encoding in ("utf-8", "gbk", "gb18030"):
        try:
            text = payload.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        return None, "这份 txt 暂时无法按常见编码读取"

    stripped = text.strip()
    if not stripped:
        return None, "这份 txt 里没有可提取的文字"
    return _limit_text(stripped)


def extract_text(payload: bytes, media_type: str, filename: str) -> tuple[str | None, str]:
    try:
        lower_name = (filename or "").lower()
        if media_type == "image":
            return ocr.image_to_text(payload, _detect_image_mime_type(payload, filename))
        if lower_name.endswith(".pdf"):
            return _extract_pdf_text(payload)
        if lower_name.endswith(".docx"):
            return _extract_docx_text(payload)
        if lower_name.endswith(".txt"):
            return _extract_txt_text(payload)
        return None, "暂不支持提取这类文件中的文字"
    except Exception as exc:  # pragma: no cover
        return None, f"提取失败({type(exc).__name__})"


def extract_image_text_with_metadata(
    payload: bytes,
    filename: str,
    *,
    evidence_id: str | None = None,
) -> tuple[str | None, str, dict[str, object]]:
    mime_type = _detect_image_mime_type(payload, filename)
    _, _, metadata = ocr._prepare_image_for_ocr_with_metadata(payload, mime_type)
    base_payload = {
        "evidence_id": evidence_id,
        "provider": "dashscope",
        "model": ocr.OCR_MODEL,
        "original_mime": metadata["original_mime"],
        "original_width": metadata["original_width"],
        "original_height": metadata["original_height"],
        "prepared_mime": metadata["prepared_mime"],
        "prepared_width": metadata["prepared_width"],
        "prepared_height": metadata["prepared_height"],
        "resized": metadata["resized"],
        "png_to_jpeg": metadata["png_to_jpeg"],
    }
    ocr._emit_image_extraction_log(
        {
            **base_payload,
            "status": "started",
            "latency_ms": 0,
            "transcript_chars": 0,
            "warning_types": [],
            "error_type": None,
        }
    )
    start = time.perf_counter()
    text, note = ocr.image_to_text(payload, mime_type)
    latency_ms = int((time.perf_counter() - start) * 1000)
    warning_type = ocr._classify_note_type(note)
    ocr._emit_image_extraction_log(
        {
            **base_payload,
            "status": "succeeded" if text is not None else "failed",
            "latency_ms": latency_ms,
            "transcript_chars": len(text) if text is not None else 0,
            "warning_types": [] if warning_type is None else [warning_type],
            "error_type": None,
        }
    )
    return text, note, metadata
